import streamlit as st
from streamlit_quill import st_quill
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
import io

# --- Cấu hình giao diện ---
st.set_page_config(page_title="Simple Word Processor", layout="wide")
st.title("Phần mềm Soạn thảo Văn bản Đơn giản")
st.markdown("Ứng dụng mô phỏng các chức năng cơ bản: định dạng văn bản và lưu file.")

# --- Hàm hỗ trợ tải file ---
def download_txt(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    text = soup.get_text()
    return io.BytesIO(text.encode("utf-8"))

def html_to_docx(html_content):
    """
    Chuyển nội dung HTML sang DOCX, giữ nguyên định dạng cơ bản.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    
    # Tìm tất cả các thẻ khối như <p> hoặc <div>
    blocks = soup.find_all(['p', 'div'])
    if not blocks:
        # Nếu không tìm thấy thẻ khối, thêm toàn bộ văn bản thuần
        doc.add_paragraph(soup.get_text())
    else:
        for block in blocks:
            paragraph = doc.add_paragraph()
            
            # Xử lý căn chỉnh từ thuộc tính style
            style = block.get('style', '')
            if 'text-align: center' in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'text-align: right' in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'text-align: justify' in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Xử lý định dạng trong đoạn
            for element in block.children:
                if element.name is None:  # Văn bản thuần túy
                    paragraph.add_run(str(element))
                elif element.name == 'strong':
                    run = paragraph.add_run(element.text)
                    run.bold = True
                elif element.name in ['em', 'i']:
                    run = paragraph.add_run(element.text)
                    run.italic = True
                elif element.name == 'u':
                    run = paragraph.add_run(element.text)
                    run.underline = True
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def html_to_pdf(html_content):
    """
    Chuyển nội dung HTML sang PDF, giữ nguyên định dạng.
    """
    result_file = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_content, dest=result_file)
    if pisa_status.err:
        st.error("Lỗi khi tạo PDF. Vui lòng kiểm tra nội dung.")
        return None
    result_file.seek(0)
    return result_file

# --- Cấu hình toolbar cho trình soạn thảo ---
toolbar = [
    ['bold', 'italic', 'underline'],  # In đậm, in nghiêng, gạch chân
    [{'align': ''}, {'align': 'center'}, {'align': 'right'}, {'align': 'justify'}],  # Căn chỉnh
]

# --- Trình soạn thảo văn bản ---
st.markdown("### Soạn thảo văn bản")
content = st_quill(
    placeholder="Nhập nội dung văn bản tại đây (hỗ trợ Ctrl+B, Ctrl+I, Ctrl+U)...",
    toolbar=toolbar,
    key="quill_editor"
)

# --- Chức năng lưu trữ ---
st.markdown("### Lưu trữ văn bản")
col1, col2, col3 = st.columns(3)

if content:
    with col1:
        txt_file = download_txt(content)
        st.download_button(
            label="Tải file TXT",
            data=txt_file,
            file_name="document.txt",
            mime="text/plain"
        )
    with col2:
        docx_file = html_to_docx(content)
        st.download_button(
            label="Tải file DOCX",
            data=docx_file,
            file_name="document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col3:
        pdf_file = html_to_pdf(content)
        if pdf_file:
            st.download_button(
                label="Tải file PDF",
                data=pdf_file,
                file_name="document.pdf",
                mime="application/pdf"
            )
else:
    st.info("Vui lòng nhập nội dung để lưu.")

# --- Hướng dẫn sử dụng ---
st.markdown("""
### Hướng dẫn sử dụng
- **Định dạng văn bản**: Sử dụng phím tắt:
  - In đậm: **Ctrl+B**
  - In nghiêng: **Ctrl+I**
  - Gạch chân: **Ctrl+U**
  - Hoặc sử dụng các nút trên thanh công cụ.
- **Căn chỉnh**: Chọn căn trái, giữa, phải, hoặc đều từ thanh công cụ.
- **Lưu trữ**: Tải file TXT, DOCX, hoặc PDF trực tiếp.
""")